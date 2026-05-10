from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\USER\Documents\SKILL--LINK\project_docs\SkillLink Project Description.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_background(section) -> None:
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        sect_pr.append(pg_borders)


def add_title(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("SkillLink")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(9, 84, 71)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Project Description and Technical Overview")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(24, 50, 79)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("A marketplace platform for connecting Zimbabwean clients with trusted local tradespeople.")
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(84, 112, 141)

    info_lines = [
        "Project Name: SkillLink",
        "Platform Focus: Clients, tradespeople, and administrators in Zimbabwe",
        "Current Build: React + TypeScript frontend, Express + TypeScript backend, Supabase schema scaffold",
    ]
    for line in info_lines:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(line)
        run.bold = True if ":" in line.split(" ", 2)[0] else False
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(38, 60, 83)

    document.add_paragraph("")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_paragraph()
    heading.style = f"Heading {level}"
    run = heading.add_run(text)
    run.font.color.rgb = RGBColor(9, 84, 71) if level == 1 else RGBColor(24, 50, 79)


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph(style="Body Text")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(38, 60, 83)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(38, 60, 83)


def add_feature_table(document: Document) -> None:
    add_heading(document, "Core Functional Scope", 1)
    add_body(
        document,
        "SkillLink currently models the main workflows needed to launch a practical services marketplace. The platform keeps the experience simple: clients can search or post requests, tradespeople can respond, and admins can monitor quality and safety.",
    )
    role_blocks = [
        (
            "Client",
            "Sees a personal dashboard, jobs page, tradespeople listings, profile, and alerts. Can post job requests, browse tradespeople, track responses, and view request activity tied to their own account.",
        ),
        (
            "Tradesperson",
            "Sees a personal workboard, jobs page, profile, nearby opportunities, and alerts. Can browse open work, apply to jobs, track their own applications, and build reputation through profile data.",
        ),
        (
            "Admin",
            "Sees an admin dashboard with requests, users, reviews, jobs, and oversight sections. Can review join requests, moderate users, inspect activity, and manage high-level platform operations.",
        ),
    ]
    for title, detail in role_blocks:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        lead = p.add_run(f"{title}: ")
        lead.bold = True
        lead.font.color.rgb = RGBColor(9, 84, 71)
        body = p.add_run(detail)
        body.font.color.rgb = RGBColor(38, 60, 83)
        body.font.size = Pt(10.5)


def build_document() -> Path:
    document = Document()

    for section in document.sections:
      section.top_margin = Inches(0.75)
      section.bottom_margin = Inches(0.75)
      section.left_margin = Inches(0.85)
      section.right_margin = Inches(0.85)
      set_page_background(section)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor(38, 60, 83)

    add_title(document)

    add_heading(document, "Project Purpose", 1)
    add_body(
        document,
        "SkillLink is designed as a Zimbabwe-first digital marketplace for household and trade services. It helps clients find reliable local workers such as plumbers, electricians, carpenters, mechanics, painters, gardeners, and other practical service providers. The core aim is to reduce friction in how people discover trusted help, especially in areas where referrals and WhatsApp already shape everyday service discovery.",
    )
    add_body(
        document,
        "The current implementation is a realistic scaffold rather than a thin mockup. It includes a themed web interface, role-aware dashboards, seeded tradespeople and location coverage, a backend API, and a relational Supabase schema prepared for production-style data handling.",
    )

    add_heading(document, "Product Vision", 1)
    add_bullets(
        document,
        [
            "Give clients a fast way to post service requests and see responses.",
            "Give tradespeople a practical channel to discover nearby work and grow visibility.",
            "Support Zimbabwean usage patterns through suburb-led search and WhatsApp-style contact expectations.",
            "Keep the platform administratively manageable with a simple but meaningful moderation layer.",
        ],
    )

    add_feature_table(document)

    add_heading(document, "Current User Experience", 1)
    add_body(
        document,
        "The frontend is built with React, TypeScript, and Vite. Authentication routes users into role-aware destinations: administrators land on the admin dashboard, while clients and tradespeople land on personal dashboards. The app includes overview, jobs, tradespeople, profile, and admin views. A protected layout wraps private pages so the navigation and role state remain consistent.",
    )
    add_bullets(
        document,
        [
            "Lightweight sign-in and account creation flow with captured contact details.",
            "Personalized client dashboard showing that client's own requests and responses.",
            "Personalized tradesperson dashboard showing that user's own applications and nearby work.",
            "Quick-request workflows so clients can enter requests with minimal friction.",
            "Clickable dashboard headings and alerts that route users toward relevant operational pages.",
        ],
    )

    add_heading(document, "Administration and Moderation", 1)
    add_body(
        document,
        "The admin role is treated as a platform operator rather than a regular marketplace user. The admin dashboard surfaces user management, join requests, jobs, reviews, reports, skills coverage, and login oversight. Approved tradesperson requests affect live counts in the dashboard model, and moderation controls are represented in the current interface for operational oversight.",
    )

    add_heading(document, "Geographic Coverage Strategy", 1)
    add_body(
        document,
        "The platform includes seeded coverage data across Zimbabwe's provinces, with a deliberate blend of CBDs, suburbs, high-density areas, and widely recognized townships such as Chitungwiza. A Zimbabwe coverage map is included in the interface to communicate supported areas visually and make location relevance obvious to users.",
    )
    add_bullets(
        document,
        [
            "Province-aware seeded coverage used in map displays and browsing context.",
            "Suburb-driven search and listing logic for both job requests and tradespeople.",
            "Location-first presentation tuned to how users commonly describe service areas in Zimbabwe.",
        ],
    )

    add_heading(document, "Technical Architecture", 1)
    add_body(
        document,
        "SkillLink uses a split frontend and backend architecture. The frontend runs as a Vite React application written in TypeScript. The backend runs as an Express API service written in TypeScript. Both services can be started from a single terminal command at the project root, which makes local development simpler for a single operator or small team.",
    )
    add_bullets(
        document,
        [
            "Frontend: React, TypeScript, React Router, local state, seeded marketplace storage, custom styling.",
            "Backend: Express, TypeScript, CORS, JSON API routes, phone utilities, OTP scaffolding, SMS integration hooks.",
            "Data layer: Supabase SQL schema with user, tradesperson, job, application, review, notification, and verification tables.",
            "Development flow: root-level scripts to run frontend and backend concurrently in one terminal.",
        ],
    )

    add_heading(document, "Database Design Summary", 1)
    add_body(
        document,
        "The Supabase schema reflects the marketplace domain clearly. It separates general users from tradesperson-specific profile details and supports the full request-application-review lifecycle. The schema also enables notifications and phone verification logs, preparing the project for stronger operational and compliance controls.",
    )

    schema_rows = [
        ("users", "Identity", "Stores role, phone, optional email, location, verification, and active status."),
        ("tradesperson_profiles", "Profile", "Adds experience, rates, availability, photos, and reputation metrics."),
        ("skills / tradesperson_skills", "Matching", "Supports searchable trade categories and many-to-many skill links."),
        ("jobs / job_applications", "Marketplace", "Tracks posted work and the tradespeople applying to each request."),
        ("reviews / notifications", "Trust + alerts", "Captures ratings and lightweight user-facing updates."),
        ("phone_verification_logs", "Operations", "Keeps a record of phone-verification delivery attempts."),
    ]
    for table_name, purpose, notes in schema_rows:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        lead = p.add_run(f"{table_name} ({purpose}): ")
        lead.bold = True
        lead.font.color.rgb = RGBColor(24, 50, 79)
        body = p.add_run(notes)
        body.font.color.rgb = RGBColor(38, 60, 83)
        body.font.size = Pt(10.5)

    add_heading(document, "Current API Surface", 1)
    add_bullets(
        document,
        [
            "GET /health for service checks.",
            "GET /api/skills for service categories.",
            "GET /api/tradespeople with skill and suburb filters.",
            "GET /api/jobs with skill and suburb filters.",
            "POST /api/auth/request-otp and POST /api/auth/verify-otp as scaffolding for phone verification flows.",
        ],
    )

    add_heading(document, "Strengths of the Current Build", 1)
    add_bullets(
        document,
        [
            "The project already expresses the business model clearly in both UI and schema form.",
            "Role-aware dashboards make the product feel more realistic than a generic demo.",
            "The seeded data set is large enough to test browsing, administration, and location-based views.",
            "The one-terminal development workflow lowers local setup friction.",
            "The codebase is ready for a gradual shift from seeded frontend state to real Supabase-backed persistence.",
        ],
    )

    add_heading(document, "Recommended Next Development Steps", 1)
    add_bullets(
        document,
        [
            "Replace local seeded account and marketplace state with Supabase persistence for multi-device consistency.",
            "Implement secure role enforcement on the backend so admin permissions are not only visual.",
            "Add production-grade phone or email authentication based on the desired launch path.",
            "Introduce proper request status updates, acceptance handling, and review submission logic.",
            "Refine moderation logging and audit trails for admin actions.",
            "Prepare deployment pipelines for the frontend, backend, and environment configuration.",
        ],
    )

    add_heading(document, "Local Development and Operations", 1)
    add_body(
        document,
        "The project is currently designed to run comfortably inside VS Code or a standard Windows terminal workflow. From the project root, running npm install followed by npm run dev starts both services together. The frontend is served on localhost:5173 and the backend on localhost:4000. Supabase configuration remains the next integration step for live data and authentication behavior.",
    )

    add_heading(document, "Conclusion", 1)
    add_body(
        document,
        "SkillLink already stands as a meaningful prototype for a Zimbabwean services marketplace. It has enough structure to explain the business model, demonstrate the user journeys, and support continued implementation without rewriting the foundation. The project now needs a focused transition from scaffolded state to persistent production data, but the current codebase gives that work a clear, well-organized starting point.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(path)
